from io import BytesIO
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from fastapi.responses import StreamingResponse
from typing import Dict, Any

class DocumentGenerator:
    @staticmethod
    def generate_word_doc(lesson_plan: Dict[str, Any]) -> BytesIO:
        doc = DocxDocument()
        
        # Add title
        doc.add_heading(lesson_plan.get('title', 'Lesson Plan'), 0)
        
        # Add basic info
        doc.add_heading('Basic Information', level=1)
        info = [
            f"School: {lesson_plan.get('school', '')}",
            f"Level: {lesson_plan.get('level', '')}",
            f"Learning Area: {lesson_plan.get('learningArea', '')}",
            f"Date: {lesson_plan.get('date', '')}",
            f"Week: {lesson_plan.get('week', '')}",
            f"Lesson: {lesson_plan.get('lessonNumber', '')}",
        ]
        
        for item in info:
            doc.add_paragraph(item)
        
        # Add learning outcomes
        if 'specificLearningOutcomes' in lesson_plan:
            doc.add_heading('Learning Outcomes', level=1)
            for outcome in lesson_plan['specificLearningOutcomes']:
                doc.add_paragraph(f"• {outcome}", style='List Bullet')
        
        # Add more sections as needed...
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_pdf(lesson_plan: Dict[str, Any]) -> BytesIO:
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Add title
        p.setFont("Helvetica-Bold", 16)
        p.drawString(72, height - 72, lesson_plan.get('title', 'Lesson Plan'))
        
        # Add basic info
        p.setFont("Helvetica", 12)
        y_position = height - 100
        info = [
            f"School: {lesson_plan.get('school', '')}",
            f"Level: {lesson_plan.get('level', '')}",
            f"Learning Area: {lesson_plan.get('learningArea', '')}",
            f"Date: {lesson_plan.get('date', '')}",
            f"Week: {lesson_plan.get('week', '')}",
            f"Lesson: {lesson_plan.get('lessonNumber', '')}",
        ]
        
        for item in info:
            p.drawString(72, y_position, item)
            y_position -= 20
        
        # Add learning outcomes
        if 'specificLearningOutcomes' in lesson_plan:
            p.setFont("Helvetica-Bold", 14)
            p.drawString(72, y_position - 30, "Learning Outcomes")
            p.setFont("Helvetica", 12)
            y_position -= 50
            
            for outcome in lesson_plan['specificLearningOutcomes']:
                p.drawString(90, y_position, f"• {outcome}")
                y_position -= 20
                if y_position < 100:
                    p.showPage()
                    y_position = height - 72
        
        p.save()
        buffer.seek(0)
        return buffer